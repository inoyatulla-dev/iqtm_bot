"""Hisobotni PDF/DOCX formatida tayyorlash.

PDF — `hisobot.html` dizayni asosida, lekin WeasyPrint uchun **0 dan** qayta yozilgan.

MUHIM: WeasyPrint `display:flex` ni to'liq qo'llab-quvvatlamaydi. A4 PDF da
sahifa chetlariga margin qo'shilganda flex-elementlar siljib, ustma-ust tushib
yoki sahifadan tashqariga chiqib ketadi. Shu sababli bu yerda BARCHA ko'p-ustunli
joylashuvlar `display:table` / `table-cell` orqali qurilgan — bu WeasyPrint da
hech qachon buzilmaydi va `@page { margin }` qiymati har qanday bo'lsa ham
kontent printable maydonga to'g'ri sig'adi (`box-sizing:border-box` + `%` enlik).
"""
import base64
import io
import math
from html import escape

from app.services.report_service import ProjectTaskRow, ReportData, StatusCount

_SUMMARY_LABELS = [
    ("Jami vazifalar", "total"),
    ("Davrda yaratilgan", "created_in_period"),
    ("Davrda bajarilgan", "done_in_period"),
    ("Kechikkan", "overdue"),
]

ORG_FULL_NAME = "Innovatsiyalarni qo'llab-quvvatlash va tijoratlashtirish markazi"

# O'zbekcha tutuq belgilari (U+02BB va h.k.) ba'zi shriftlarda ▀ bo'lib chiqadi —
# oddiy ASCII apostrofga almashtiramiz.
_APOSTROPHE_VARIANTS = "ʻʼ‘’´`′"


# ── Yordamchi funksiyalar ───────────────────────────────────────────────────

def _clean(value) -> str:
    text = "" if value is None else str(value)
    for ch in _APOSTROPHE_VARIANTS:
        text = text.replace(ch, "'")
    return text


def _esc(value) -> str:
    return escape(_clean(value), quote=False)


def _summary_rows(data: ReportData) -> list[tuple[str, int]]:
    return [(label, getattr(data, attr)) for label, attr in _SUMMARY_LABELS]


def _tint(hex_color: str, alpha: float = 0.14) -> str:
    """Holat rangidan ochiq fon (badge foni uchun) hosil qiladi."""
    hex_color = (hex_color or "").lstrip("#")
    if len(hex_color) == 3:
        hex_color = "".join(ch * 2 for ch in hex_color)
    try:
        r, g, b = (int(hex_color[i:i + 2], 16) for i in (0, 2, 4))
    except ValueError:
        r, g, b = 0x64, 0x74, 0x8B
    return f"rgba({r}, {g}, {b}, {alpha})"


def _bar_width(percent: int, total: int) -> str:
    """Progress-bar to'ldirilishi — 0% bo'lsa bo'sh, aks holda ko'rinish uchun min. 3%."""
    if not total or percent <= 0:
        return "0%"
    return f"{max(percent, 3)}%"


def _progress_color(percent: int) -> str:
    if percent >= 100:
        return "#16A34A"   # yashil — to'liq
    if percent >= 50:
        return "#2563EB"   # ko'k — yarmidan ko'p
    if percent > 0:
        return "#D97706"   # sariq — past
    return "#CBD2D8"       # kulrang — 0%


def _kpi_sub(data: ReportData, attr: str) -> str:
    if attr == "total":
        return "joriy faol"
    if attr == "created_in_period":
        return "shu davrda"
    if attr == "done_in_period":
        percent = round(data.done_in_period * 100 / data.total) if data.total else 0
        return f"{percent}% bajarilish"
    if attr == "overdue":
        return "e'tibor talab etadi"
    return ""


# ── Donut (SVG — conic-gradient WeasyPrint da ishlamaydi) ────────────────────

def _donut_svg(statuses: list[StatusCount], total: int, size: int = 156) -> str:
    """Holatlar taqsimoti — to'ldirilgan doiraviy diagramma + markazda umumiy son."""
    cx = cy = size / 2
    r = size / 2
    hole_r = r * 0.62

    nonzero = [s for s in statuses if s.count > 0]
    if total <= 0:
        pie = f'<circle cx="{cx}" cy="{cy}" r="{r}" fill="#EFF2F4"/>'
    elif len(nonzero) == 1:
        pie = f'<circle cx="{cx}" cy="{cy}" r="{r}" fill="{nonzero[0].color}"/>'
    else:
        segments = []
        angle = -90.0
        for s in nonzero:
            sweep = s.count / total * 360
            a0, a1 = math.radians(angle), math.radians(angle + sweep)
            x1, y1 = cx + r * math.cos(a0), cy + r * math.sin(a0)
            x2, y2 = cx + r * math.cos(a1), cy + r * math.sin(a1)
            large = 1 if sweep > 180 else 0
            segments.append(
                f'<path d="M{cx},{cy} L{x1:.3f},{y1:.3f} '
                f'A{r},{r} 0 {large} 1 {x2:.3f},{y2:.3f} Z" fill="{s.color}"/>'
            )
            angle += sweep
        pie = "".join(segments)

    hole = f'<circle cx="{cx}" cy="{cy}" r="{hole_r}" fill="#fff"/>'
    n_size = hole_r * 0.52
    l_size = hole_r * 0.18
    text = (
        f'<text x="{cx}" y="{cy + n_size * 0.30:.1f}" text-anchor="middle" '
        f'font-family="\'Plus Jakarta Sans\',sans-serif" font-size="{n_size:.1f}" '
        f'font-weight="800" fill="#111827">{total}</text>'
        f'<text x="{cx}" y="{cy + n_size * 0.30 + l_size + 9:.1f}" text-anchor="middle" '
        f'font-family="Inter,sans-serif" font-size="{l_size:.1f}" font-weight="600" '
        f'letter-spacing="1" fill="#6B7280">VAZIFA</text>'
    )
    return (
        f'<svg width="{size}" height="{size}" viewBox="0 0 {size} {size}" '
        f'xmlns="http://www.w3.org/2000/svg">{pie}{hole}{text}</svg>'
    )


# ── CSS (faqat WeasyPrint qo'llab-quvvatlaydigan xususiyatlar) ───────────────

_HTML_CSS = """
@page {
  size: A4;
  margin: 12mm 12mm 14mm 12mm;
  @bottom-left {
    content: "IQTM hisoboti";
    font-family: 'Inter', sans-serif; font-size: 8pt; color: #9AA0A6;
  }
  @bottom-right {
    content: "Sahifa " counter(page) " / " counter(pages);
    font-family: 'Inter', sans-serif; font-size: 8pt; color: #9AA0A6;
  }
}

* { box-sizing: border-box; margin: 0; padding: 0; }

body {
  font-family: 'Inter', -apple-system, "Segoe UI", Roboto, sans-serif;
  color: #111827; background: #fff;
  font-size: 12px; line-height: 1.45;
}
.sheet { width: 100%; }
.pad { padding: 20px 24px 6px; }

h2.section {
  font-family: 'Plus Jakarta Sans', sans-serif;
  font-size: 12px; font-weight: 700; letter-spacing: .06em;
  text-transform: uppercase; color: #0A5B42;
  margin: 22px 0 12px; break-after: avoid;
}
h2.section .bar {
  display: inline-block; width: 5px; height: 13px; border-radius: 3px;
  background: #0E7C5A; vertical-align: -2px; margin-right: 9px;
}
h2.section:first-child { margin-top: 0; }

/* ---------- HEADER ---------- */
.header {
  background: linear-gradient(135deg, #0A5B42, #0E7C5A);
  color: #fff; padding: 22px 24px;
}
.header__row { display: table; width: 100%; }
.header__left { display: table-cell; vertical-align: top; }
.header__right {
  display: table-cell; vertical-align: top; text-align: right; white-space: nowrap;
}
.header__brand { display: table; }
.header__logo-cell { display: table-cell; vertical-align: middle; width: 44px; }
.header__name-cell { display: table-cell; vertical-align: middle; padding-left: 12px; }
.header__logo {
  width: 44px; height: 44px; border-radius: 10px; background: #fff;
  color: #0E7C5A; font-family: 'Plus Jakarta Sans', sans-serif; font-weight: 800;
  font-size: 13px; text-align: center; line-height: 44px;
}
img.header__logo { object-fit: contain; }
.header__name { font-size: 12px; font-weight: 600; opacity: .92; line-height: 1.3; max-width: 300px; }
.header__doctype {
  font-family: 'Plus Jakarta Sans', sans-serif; font-size: 10px; font-weight: 700;
  letter-spacing: .14em; text-transform: uppercase; opacity: .8; margin-top: 16px;
}
.header__h1 {
  font-family: 'Plus Jakarta Sans', sans-serif; font-size: 26px; font-weight: 800;
  letter-spacing: -.02em; margin-top: 2px;
}
.header__meta .period { font-weight: 700; font-size: 14px; }
.header__meta .gen { margin-top: 6px; opacity: .85; font-size: 12px; }

/* ---------- KPI CARDS (table, border-spacing bilan gap) ---------- */
.kpis {
  display: table; width: 100%; table-layout: fixed;
  border-spacing: 10px 0; margin: 0 -10px;
}
.kpi {
  display: table-cell; width: 25%; vertical-align: top; background: #fff;
  border: 1px solid #E6E8EB; border-left: 4px solid #0E7C5A; border-radius: 10px;
  padding: 13px 15px;
}
.kpi.accent-red { border-left-color: #DC2626; }
.kpi .label {
  font-size: 10px; font-weight: 600; color: #6B7280;
  text-transform: uppercase; letter-spacing: .04em;
}
.kpi .value {
  font-family: 'Plus Jakarta Sans', sans-serif; font-size: 24px; font-weight: 800;
  line-height: 1.1; margin-top: 5px;
}
.kpi.accent-red .value { color: #DC2626; }
.kpi .sub { font-size: 10.5px; color: #6B7280; margin-top: 2px; }

/* ---------- STATUS: donut + legend (table) ---------- */
.status-wrap {
  display: table; width: 100%; background: #fff;
  border: 1px solid #E6E8EB; border-radius: 10px; padding: 16px 20px;
  break-inside: avoid;
}
.status-wrap__donut { display: table-cell; width: 168px; vertical-align: middle; }
.status-wrap__legend { display: table-cell; vertical-align: middle; padding-left: 24px; }
.legend { display: table; width: 100%; }
.legend__row { display: table-row; }
.legend__cell { display: table-cell; width: 50%; padding: 5px 16px 5px 0; vertical-align: middle; }
.leg { display: table; width: 100%; }
.leg.zero { opacity: .45; }
.leg .dot { display: table-cell; width: 16px; vertical-align: middle; }
.leg .dot i { display: inline-block; width: 10px; height: 10px; border-radius: 3px; }
.leg .ln { display: table-cell; vertical-align: middle; font-size: 12px; font-weight: 500; color: #111827; }
.leg .cnt {
  display: table-cell; vertical-align: middle; text-align: right; width: 34px;
  font-family: 'Plus Jakarta Sans', sans-serif; font-weight: 700; font-size: 12px;
}

/* ---------- DEPARTMENTS: progress bars (table) ---------- */
.depts { display: table; width: 100%; }
.dept { display: table-row; }
.dept > div { display: table-cell; vertical-align: middle; padding: 8px 0; border-bottom: 1px solid #E6E8EB; }
.dept:last-child > div { border-bottom: none; }
.dept .nm { width: 130px; font-weight: 600; font-size: 12px; padding-right: 14px; }
.dept .track-cell { padding-right: 14px; }
.dept .track { width: 100%; height: 8px; border-radius: 6px; background: #EFF2F4; overflow: hidden; }
.dept .fill { height: 8px; border-radius: 6px; }
.dept .pct { width: 96px; text-align: right; font-size: 11px; color: #6B7280; }
.dept .pct b { font-family: 'Plus Jakarta Sans', sans-serif; color: #111827; font-size: 12px; }

/* ---------- PROJECTS ---------- */
.project { border: 1px solid #E6E8EB; border-radius: 10px; margin-bottom: 14px; background: #fff; }
.project.late { border-color: #F1B0B0; }
.proj-head {
  padding: 12px 16px; background: linear-gradient(180deg, #FAFBFC, #fff);
  border-bottom: 1px solid #E6E8EB; border-radius: 10px 10px 0 0; break-after: avoid;
}
.project.late .proj-head { background: linear-gradient(180deg, #FEF4F4, #fff); border-bottom-color: #F1B0B0; }
.proj-head__row { display: table; width: 100%; }
.proj-title { display: table-cell; vertical-align: middle; }
.proj-title .pn { font-family: 'Plus Jakarta Sans', sans-serif; font-size: 14px; font-weight: 700; }
.pill {
  display: inline-block; font-size: 10px; font-weight: 700; letter-spacing: .04em;
  text-transform: uppercase; padding: 3px 9px; border-radius: 20px;
  background: #E7F4EF; color: #0A5B42; margin-left: 10px; vertical-align: middle;
}
.pill.late { background: #FEE2E2; color: #DC2626; }
.proj-prog { display: table-cell; vertical-align: middle; text-align: right; width: 200px; white-space: nowrap; }
.proj-prog .ptrack {
  display: inline-block; width: 104px; height: 7px; border-radius: 6px;
  background: #EFF2F4; overflow: hidden; vertical-align: middle;
}
.proj-prog .pfill { height: 7px; border-radius: 6px; }
.proj-prog .ptxt { display: inline-block; font-size: 12px; color: #6B7280; margin-left: 10px; vertical-align: middle; }
.proj-prog .ptxt b { font-family: 'Plus Jakarta Sans', sans-serif; color: #111827; }
.proj-sub { padding: 8px 16px; font-size: 12px; color: #6B7280; border-bottom: 1px solid #E6E8EB; }
.proj-sub span { margin-right: 20px; }
.proj-sub b { color: #111827; font-weight: 600; }
.proj-sub .warn { color: #DC2626; font-weight: 600; }

/* ---------- TASK TABLE (native, auto layout — overflow'siz) ---------- */
table.tasks { width: 100%; border-collapse: collapse; }
table.tasks thead { display: table-header-group; }   /* sahifa bo'linsa header takrorlanadi */
table.tasks th {
  font-size: 10px; font-weight: 700; letter-spacing: .05em; text-transform: uppercase;
  color: #6B7280; text-align: left; padding: 8px 16px; background: #F8FAFB;
  border-bottom: 1px solid #E6E8EB;
}
table.tasks td { padding: 9px 16px; font-size: 12px; border-bottom: 1px solid #F0F2F4; vertical-align: middle; }
table.tasks tr { break-inside: avoid; }
table.tasks tbody tr:last-child td { border-bottom: none; }
td.num, th.c-num { width: 26px; text-align: center; }
td.num { color: #6B7280; font-family: 'Plus Jakarta Sans', sans-serif; font-weight: 700; }
th.c-who, td.who { width: 130px; }
td.who { color: #111827; font-weight: 500; }
th.c-due, td.due { width: 96px; white-space: nowrap; }
td.due { color: #6B7280; }
td.due.over { color: #DC2626; font-weight: 600; }
th.c-status, td.status { width: 104px; }
.badge {
  display: inline-block; font-size: 11px; font-weight: 600;
  padding: 4px 10px; border-radius: 20px; white-space: nowrap;
}
.badge i {
  display: inline-block; width: 6px; height: 6px; border-radius: 50%;
  background: currentColor; margin-right: 6px; vertical-align: middle;
}

.empty { color: #6B7280; font-size: 12px; padding: 8px 16px; }
"""


# ── HTML bo'laklar generatori ───────────────────────────────────────────────

def _kpi_html(data: ReportData) -> str:
    return "".join(
        f'<div class="kpi{" accent-red" if attr == "overdue" else ""}">'
        f'<div class="label">{_esc(label)}</div>'
        f'<div class="value">{getattr(data, attr)}</div>'
        f'<div class="sub">{_esc(_kpi_sub(data, attr))}</div></div>'
        for label, attr in _SUMMARY_LABELS
    )


def _legend_html(statuses: list[StatusCount]) -> str:
    """Holatlarni 2 ustunli jadval ko'rinishida (har biri: nuqta + nom + son)."""
    legs = [
        f'<div class="legend__cell"><div class="leg{" zero" if s.count == 0 else ""}">'
        f'<div class="dot"><i style="background:{s.color}"></i></div>'
        f'<div class="ln">{_esc(s.name)}</div>'
        f'<div class="cnt">{s.count}</div></div></div>'
        for s in statuses
    ]
    if not legs:
        return '<div class="empty">Holatlar mavjud emas</div>'
    rows = ""
    for i in range(0, len(legs), 2):
        pair = legs[i:i + 2]
        if len(pair) == 1:
            pair.append('<div class="legend__cell"></div>')  # juftlik uchun bo'sh katak
        rows += f'<div class="legend__row">{"".join(pair)}</div>'
    return rows


def _depts_html(data: ReportData) -> str:
    depts_sorted = sorted(data.departments, key=lambda d: d.percent, reverse=True)
    rows = "".join(
        f'<div class="dept">'
        f'<div class="nm">{_esc(d.name)}</div>'
        f'<div class="track-cell"><div class="track">'
        f'<div class="fill" style="width:{_bar_width(d.percent, d.total)};'
        f'background:{_progress_color(d.percent)}"></div></div></div>'
        f'<div class="pct"><b>{f"{d.percent}%" if d.total else "—"}</b> · {d.done}/{d.total}</div>'
        f"</div>"
        for d in depts_sorted
    )
    return rows or '<div class="empty">Bo\'limlar mavjud emas</div>'


def _task_row(t: ProjectTaskRow) -> str:
    if t.overdue:
        bg, color = "rgba(220, 38, 38, 0.12)", "#DC2626"
    else:
        bg, color = _tint(t.status_color), t.status_color
    due_cls = "due over" if t.overdue else "due"
    return (
        f'<tr>'
        f'<td class="num">{t.seq}</td>'
        f'<td class="task">{_esc(t.name)}</td>'
        f'<td class="who">{_esc(t.assignee)}</td>'
        f'<td class="{due_cls}">{_esc(t.deadline)}</td>'
        f'<td class="status"><span class="badge" style="background:{bg};color:{color}">'
        f'<i></i>{_esc(t.status_label)}</span></td>'
        f"</tr>"
    )


def _projects_html(data: ReportData) -> str:
    blocks = ""
    for proj in data.projects:
        late = proj.schedule_label.startswith("Orqada qolmoqda") or \
            proj.schedule_label.startswith("Muddat o'tgan")
        pill_label = "Orqada qolmoqda" if late else proj.status_label
        width = _bar_width(proj.percent, proj.total)
        pfill = "#DC2626" if late else _progress_color(proj.percent)
        rows = "".join(_task_row(t) for t in proj.tasks) or \
            '<tr><td colspan="5" class="empty">Vazifalar mavjud emas</td></tr>'
        second = (
            f'<span class="warn">⚠ {_esc(proj.schedule_label)}</span>'
            if late else
            f'<span>Joriy jarayon: <b>{_esc(proj.schedule_label)}</b></span>'
        )
        blocks += (
            f'<div class="project{" late" if late else ""}">'
            f'<div class="proj-head"><div class="proj-head__row">'
            f'<div class="proj-title"><span class="pn">{_esc(proj.name)}</span>'
            f'<span class="pill{" late" if late else ""}">{_esc(pill_label)}</span></div>'
            f'<div class="proj-prog"><span class="ptrack"><span class="pfill" '
            f'style="display:block;width:{width};background:{pfill}"></span></span>'
            f'<span class="ptxt"><b>{proj.done}/{proj.total}</b> · {proj.percent}%</span></div>'
            f'</div></div>'
            f'<div class="proj-sub"><span>Muddat: <b>{_esc(proj.deadline_label)}</b></span>{second}</div>'
            f'<table class="tasks"><thead><tr>'
            f'<th class="c-num">#</th><th class="c-task">Vazifa</th>'
            f'<th class="c-who">Mas\'ul</th><th class="c-due">Muddat</th>'
            f'<th class="c-status">Holat</th></tr></thead>'
            f'<tbody>{rows}</tbody></table>'
            f"</div>"
        )
    return blocks or '<div class="empty">Loyihalar mavjud emas</div>'


def _logo_html(data: ReportData) -> str:
    if data.logo_path and data.logo_path.exists():
        b64 = base64.b64encode(data.logo_path.read_bytes()).decode("ascii")
        return f'<img class="header__logo" src="data:image/png;base64,{b64}" alt="IQTM">'
    return '<div class="header__logo">IQTM</div>'


def build_html(data: ReportData) -> str:
    """`hisobot.html` ko'rinishidagi to'liq hisobot — A4 PDF ga tayyor, WeasyPrint-safe."""
    org_name = _clean(data.org_name) or ORG_FULL_NAME
    total_status = sum(s.count for s in data.statuses)

    return f"""<!DOCTYPE html>
<html lang="uz">
<head>
<meta charset="UTF-8">
<title>{_esc(data.period_title)} hisobot</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@500;600;700;800&family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>{_HTML_CSS}</style>
</head>
<body>
<div class="sheet">
  <div class="header"><div class="header__row">
    <div class="header__left">
      <div class="header__brand">
        <div class="header__logo-cell">{_logo_html(data)}</div>
        <div class="header__name-cell"><div class="header__name">{_esc(org_name)}</div></div>
      </div>
      <div class="header__doctype">{_esc(data.period_title)} hisobot</div>
      <div class="header__h1">Hisobot</div>
    </div>
    <div class="header__right header__meta">
      <div class="period">{data.start.strftime('%d.%m.%Y')} — {data.end.strftime('%d.%m.%Y')}</div>
      <div class="gen">Yaratildi: {data.generated_at.strftime('%d.%m.%Y, %H:%M')}</div>
    </div>
  </div></div>

  <div class="pad">
    <h2 class="section"><span class="bar"></span>Umumiy ko'rsatkichlar</h2>
    <div class="kpis">{_kpi_html(data)}</div>

    <h2 class="section"><span class="bar"></span>Holatlar bo'yicha</h2>
    <div class="status-wrap">
      <div class="status-wrap__donut">{_donut_svg(data.statuses, total_status)}</div>
      <div class="status-wrap__legend"><div class="legend">{_legend_html(data.statuses)}</div></div>
    </div>

    <h2 class="section"><span class="bar"></span>Bo'limlar bo'yicha</h2>
    <div class="depts">{_depts_html(data)}</div>

    <h2 class="section"><span class="bar"></span>Loyihalar</h2>
    <div class="loyihalar">{_projects_html(data)}</div>
  </div>
</div>
</body>
</html>"""


def build_pdf(data: ReportData) -> bytes:
    """HTML hisobotni A4 PDF'ga aylantiradi (WeasyPrint)."""
    from weasyprint import HTML

    return HTML(string=build_html(data)).write_pdf()


# ── DOCX (oldingi ishlaydigan versiya — o'zgartirilmagan) ────────────────────

def build_docx(data: ReportData) -> bytes:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    navy = RGBColor(0x1E, 0x3A, 0x5F)
    accent = RGBColor(0x25, 0x63, 0xEB)
    org_name = _clean(data.org_name) or ORG_FULL_NAME

    doc = Document()

    top_table = doc.add_table(rows=1, cols=2)
    top_table.autofit = False
    top_table.allow_autofit = False
    top_table.columns[0].width = Inches(0.6)
    top_table.columns[1].width = Inches(5.9)
    logo_cell, org_cell = top_table.rows[0].cells
    logo_cell.width = Inches(0.6)
    org_cell.width = Inches(5.9)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    org_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if data.logo_path and data.logo_path.exists():
        logo_cell.paragraphs[0].add_run().add_picture(str(data.logo_path), width=Inches(0.45))
    run = org_cell.paragraphs[0].add_run(org_name)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = navy

    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = line_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_heading("Hisobot", level=0)
    doc.add_paragraph(f"Davr: {data.start} — {data.end}")
    doc.add_paragraph(f"Yaratildi: {data.generated_at.strftime('%Y-%m-%d %H:%M')}")

    def section_title(text: str, color: RGBColor = navy):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = color

    def fill_table(headers, rows, widths=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for r in hdr[i].paragraphs[0].runs:
                r.font.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = _clean(v)
        if widths:
            for row in table.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Inches(w)
        return table

    section_title("Umumiy ko'rsatkichlar")
    fill_table(["Ko'rsatkich", "Qiymat"],
               [[label, str(value)] for label, value in _summary_rows(data)],
               widths=[4.2, 1.8])

    if data.statuses:
        section_title("Holatlar bo'yicha")
        fill_table(["Holat", "Soni"],
                   [[s.name, str(s.count)] for s in data.statuses],
                   widths=[4.2, 1.8])

    if data.departments:
        section_title("Bo'limlar bo'yicha")
        fill_table(["Bo'lim", "Jami", "Bajarilgan", "Foiz"],
                   [[d.name, str(d.total), str(d.done), f"{d.percent}%"] for d in data.departments],
                   widths=[3, 1, 1, 1])

    if data.projects:
        section_title("Loyihalar")
        for proj in data.projects:
            h = doc.add_heading(_clean(proj.name), level=3)
            for r in h.runs:
                r.font.color.rgb = accent
            fill_table(["", ""],
                       [["Holat", proj.status_label],
                        ["Muddat", proj.deadline_label],
                        ["Bajarilgan vazifalar", f"{proj.done}/{proj.total} ({proj.percent}%)"],
                        ["Joriy jarayon", proj.schedule_label]],
                       widths=[2, 4])
            doc.add_paragraph()
            if proj.tasks:
                fill_table(["#", "Vazifa", "Mas'ul", "Muddat", "Holat"],
                           [[t.seq, t.name, t.assignee, t.deadline, t.status_label] for t in proj.tasks],
                           widths=[0.4, 2.6, 1.6, 1.2, 1.2])
            else:
                doc.add_paragraph("Vazifalar mavjud emas")
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()